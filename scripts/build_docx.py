"""把 docs/中期汇报.md 转成 docx 交付物。"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "中期汇报.md"
OUT = ROOT / "reports" / "视觉质量评估_中期汇报.docx"


def _set_font(run, name="PingFang SC", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    if rFonts is None:
        from docx.oxml.ns import qn
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    from docx.oxml.ns import qn
    rFonts.set(qn("w:eastAsia"), name)


def _add_paragraph(doc, text, style=None, size=11, bold=False, color=None,
                   alignment=None):
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return p


def md_to_docx(src: Path, out: Path):
    text = src.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style.font.size = Pt(11)

    in_table = False
    table_rows: list[list[str]] = []

    in_code = False
    code_buffer: list[str] = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        # 跳过分隔行 |---|---|
        clean = [r for r in table_rows if not re.match(r"^\s*\|?[-:\s|]+\|?\s*$", "|".join(r))]
        if not clean:
            table_rows = []
            return
        ncol = max(len(r) for r in clean)
        tbl = doc.add_table(rows=len(clean), cols=ncol)
        tbl.style = "Light Grid Accent 1"
        for i, row in enumerate(clean):
            for j in range(ncol):
                cell_text = row[j] if j < len(row) else ""
                cell_text = cell_text.replace("**", "")
                tcell = tbl.rows[i].cells[j]
                tcell.text = ""
                p = tcell.paragraphs[0]
                run = p.add_run(cell_text.strip())
                _set_font(run, size=10, bold=(i == 0))
        doc.add_paragraph()
        table_rows = []

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buffer))
        _set_font(run, name="Menlo", size=9, color=(50, 80, 140))
        # 浅灰底
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F1F5F9")
        pPr.append(shd)
        code_buffer = []

    for raw in lines:
        line = raw.rstrip()

        # code fence
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue

        # Tables
        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            in_table = True
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        if not line.strip():
            doc.add_paragraph()
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            content = m.group(2)
            sizes = {1: 22, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}
            _add_paragraph(doc, content,
                           size=sizes[level], bold=True,
                           color=(15, 23, 42) if level == 1 else (37, 99, 235))
            continue

        # Quote
        if line.startswith("> "):
            _add_paragraph(doc, line[2:], size=10, color=(100, 116, 139))
            continue

        # bullet
        if re.match(r"^\s*[-*]\s+", line):
            txt = re.sub(r"^\s*[-*]\s+", "", line)
            txt = txt.replace("**", "")
            doc.add_paragraph(txt, style="List Bullet")
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            txt = re.sub(r"^\s*\d+\.\s+", "", line)
            txt = txt.replace("**", "")
            doc.add_paragraph(txt, style="List Number")
            continue

        # Plain paragraph: handle bold **xxx**
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", line)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                _set_font(run, bold=True)
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                _set_font(run, name="Menlo", size=10, color=(37, 99, 235))
            else:
                run = p.add_run(part)
                _set_font(run)

    flush_table()
    flush_code()

    # 在第二节的某处插入截图
    doc.add_page_break()
    _add_paragraph(doc, "附录：GUI 主界面截图", size=14, bold=True,
                   color=(37, 99, 235))
    pic_path = ROOT / "reports" / "gui_screenshot.png"
    if pic_path.exists():
        doc.add_picture(str(pic_path), width=Cm(16))

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"saved: {out}")


if __name__ == "__main__":
    md_to_docx(SRC, OUT)
